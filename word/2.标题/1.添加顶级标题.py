"""
@from：https://pythoneers.cn
@author: qq3330447288
@contact: erics1996@yeah.net
@software: PyCharm
@file: 添加标题.py
@time: 2020/11/15 下午12:21
"""
from docx import Document

document = Document()
document.add_heading('The REAL meaning of the universe')
document.add_paragraph('1. Lorem ipsum dolor sit amet.')
document.add_paragraph('2. Lorem ipsum dolor sit amet.')
document.save('./result/添加标题-顶级标题.docx')