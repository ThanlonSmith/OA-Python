"""
@from：https://pythoneers.cn
@author: qq3330447288
@contact: erics1996@yeah.net
@software: PyCharm
@file: 添加小结标题.py
@time: 2020/11/15 下午12:26
"""
from docx import Document

document = Document()
document.add_heading('The REAL meaning of the universe', level=2)
document.add_paragraph('1. Lorem ipsum dolor sit amet.')
document.add_heading('The REAL meaning of the universe', level=3)
document.add_paragraph('1. Lorem ipsum dolor sit amet.')
document.add_heading('The REAL meaning of the universe', level=4)
document.add_paragraph('1. Lorem ipsum dolor sit amet.')
document.add_heading('The REAL meaning of the universe', level=0)
document.add_paragraph('1. Lorem ipsum dolor sit amet.')
document.save('./result/添加标题-小结标题.docx')
