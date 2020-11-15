"""
@from：https://pythoneers.cn
@author: qq3330447288
@contact: erics1996@yeah.net
@software: PyCharm
@file: 应用段落样式.py
@time: 2020/11/15 下午6:35
"""
from docx import Document

document = Document()
# document.add_paragraph('Lorem ipsum dolor sit amet.', style='ListBullet')
paragraph = document.add_paragraph('Lorem ipsum dolor sit amet.')
paragraph.style = 'ListBullet'
document.save('./result/应用段落样式.docx')
