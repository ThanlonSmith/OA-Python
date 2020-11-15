"""
@from：https://pythoneers.cn
@author: qq3330447288
@contact: erics1996@yeah.net
@software: PyCharm
@file: 添加分页.py
@time: 2020/11/15 下午1:07
"""
from docx import Document

document = Document()
document.add_paragraph('1. Lorem ipsum dolor sit amet.')
document.add_page_break()
document.add_paragraph('2. Lorem ipsum dolor sit amet.')
document.save('./result/添加分页.docx')
