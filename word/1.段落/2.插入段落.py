"""
@from：https://pythoneers.cn
@author: qq3330447288
@contact: erics1996@yeah.net
@software: PyCharm
@file: 插入段落.py
@time: 2020/11/15 上午11:24
"""
from docx import Document

document = Document()
document.add_paragraph('1. Lorem ipsum dolor sit amet.')
document.add_paragraph('2. Lorem ipsum dolor sit amet.')
paragraph = document.add_paragraph('3. Lorem ipsum dolor sit amet.')
paragraph.insert_paragraph_before('4. Lorem ipsum dolor sit amet.')
document.save('./result/插入段落.docx')
