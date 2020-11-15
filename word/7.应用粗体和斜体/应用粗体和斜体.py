"""
@from：https://pythoneers.cn
@author: qq3330447288
@contact: erics1996@yeah.net
@software: PyCharm
@file: 应用粗体和斜体.py
@time: 2020/11/15 下午7:33
"""
from docx import Document

document = Document()
paragraph = document.add_paragraph()
paragraph.add_run('Lorem ipsum ')
paragraph.add_run('dolor').bold = True
"""
# paragraph.add_run('dolor').bold = True等价于
run = paragraph.add_run('dolor')
run.bold = True
"""
paragraph.add_run(' sit amet.')
document.save('./result/应用粗体和斜体.docx')
