"""
@from：https://pythoneers.cn
@author: qq3330447288
@contact: erics1996@yeah.net
@software: PyCharm
@file: 应用字体样式.py
@time: 2020/11/15 下午7:47
"""
from docx import Document

document = Document()
paragraph = document.add_paragraph('Normal text, ')
paragraph.add_run('text with emphasis.', 'Emphasis')
document.save('./result/应用字体样式.docx')
